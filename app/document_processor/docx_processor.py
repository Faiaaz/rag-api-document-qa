"""
Word document processor using python-docx
"""
from docx import Document
from typing import List, Dict, Any
from pathlib import Path

from .base_processor import BaseDocumentProcessor


class DOCXProcessor(BaseDocumentProcessor):
    """Processor for Word documents (.docx)"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size, chunk_overlap)
        self.supported_extensions = {'.docx', '.doc'}
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the given file"""
        if not self.validate_file(file_path):
            return False
        
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted text content
        """
        try:
            # Open the document
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("\n".join(table_text))
            
            # Join all text content
            full_text = "\n\n".join(text_content)
            
            return full_text
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from Word document {file_path}: {str(e)}")
    
    def extract_text_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text with additional metadata from Word document
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Get document properties
            core_props = doc.core_properties
            
            # Extract text
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("\n".join(table_text))
            
            return {
                "text": "\n\n".join(text_content),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            raise ValueError(f"Failed to extract text and metadata from Word document {file_path}: {str(e)}") 